"""Tests for career asset loaders (docx, md, txt) and normalization.

Uses synthetic files in tmp_path. Never touches a real career document, the
network, or the database.
"""

import pytest
from docx import Document

from app.career.assets import CareerAssetType
from app.career.loaders import (
    load_career_asset,
    load_docx_asset,
    load_text_asset,
    normalize_asset_text,
)

BRIEF = CareerAssetType.MASTER_CAREER_BRIEF


def _long_text(seed: str) -> str:
    return (
        f"{seed} paragraph with enough words to comfortably exceed the "
        "one-hundred character minimum used by the loaders, plus some extra "
        "content so the assertion on wording preservation is meaningful."
    )


def _write_docx(path, paragraphs, table=None):
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table is not None:
        rows = len(table)
        cols = len(table[0])
        doc_table = doc.add_table(rows=rows, cols=cols)
        for i, row in enumerate(table):
            for j, value in enumerate(row):
                doc_table.cell(i, j).text = value
    doc.save(str(path))


def test_docx_paragraphs_are_extracted(tmp_path):
    text = _long_text("First paragraph")
    path = tmp_path / "brief.docx"
    _write_docx(path, [text])

    asset = load_docx_asset(path, BRIEF)

    assert text in asset.text
    assert asset.source_format == "docx"
    assert asset.character_count == len(asset.text)


def test_docx_table_cells_are_extracted(tmp_path):
    text = _long_text("Table cell content")
    path = tmp_path / "brief.docx"
    _write_docx(path, ["Intro"], table=[["Header", "Value"], ["cell A", text]])

    asset = load_docx_asset(path, BRIEF)

    assert "cell A" in asset.text
    assert text in asset.text


def test_docx_source_order_is_preserved(tmp_path):
    first = _long_text("Paragraph one")
    last = _long_text("Paragraph three")
    path = tmp_path / "brief.docx"
    _write_docx(path, [first, "Para two"], table=[["MIDDLE_CELL"]] + [[last]])

    asset = load_docx_asset(path, BRIEF)

    assert asset.text.index(first) < asset.text.index("MIDDLE_CELL")
    assert asset.text.index("MIDDLE_CELL") < asset.text.index(last)


def test_markdown_and_text_files_are_extracted(tmp_path):
    for suffix, fmt in ((".md", "md"), (".txt", "txt")):
        text = _long_text(f"content {suffix}")
        path = tmp_path / f"asset{suffix}"
        path.write_text(text, encoding="utf-8")

        asset = load_text_asset(path, BRIEF)

        assert text in asset.text
        assert asset.source_format == fmt


def test_line_endings_are_normalized(tmp_path):
    path = tmp_path / "asset.txt"
    text = (
        _long_text("A")
        + "\r\n"
        + _long_text("B")
        + "\r"
        + _long_text("C")
    )
    path.write_text(text, encoding="utf-8")

    asset = load_text_asset(path, BRIEF)

    assert "\r" not in asset.text
    assert "\n" in asset.text


def test_blank_line_runs_are_collapsed(tmp_path):
    path = tmp_path / "asset.txt"
    text = (
        _long_text("First line")
        + "\n\n\n\n\n"
        + _long_text("Second line")
    )
    path.write_text(text, encoding="utf-8")

    asset = load_text_asset(path, BRIEF)

    # A run of four blank lines collapses to two (rendered as "\n\n\n");
    # a run of three or more blank lines must no longer be present.
    assert "\n\n\n" in asset.text
    assert "\n\n\n\n" not in asset.text


def test_hashes_are_deterministic(tmp_path):
    path = tmp_path / "asset.txt"
    path.write_text(_long_text("stable"), encoding="utf-8")

    first = load_text_asset(path, BRIEF)
    second = load_text_asset(path, BRIEF)

    assert first.content_hash == second.content_hash
    assert len(first.content_hash) == 64


def test_missing_file_fails_clearly(tmp_path):
    with pytest.raises(FileNotFoundError):
        load_docx_asset(tmp_path / "missing.docx", BRIEF)


def test_unsupported_format_fails_clearly(tmp_path):
    path = tmp_path / "asset.pdf"
    path.write_text(_long_text("pdf-like"), encoding="utf-8")

    with pytest.raises(ValueError):
        load_career_asset(path, BRIEF)
    with pytest.raises(ValueError):
        load_text_asset(path, BRIEF)


def test_full_source_wording_is_preserved(tmp_path):
    text = _long_text("Distinctive wording")
    path = tmp_path / "brief.docx"
    _write_docx(path, [text])

    asset = load_docx_asset(path, BRIEF)

    assert asset.text == normalize_asset_text(text)


def test_repr_does_not_expose_full_asset_text(tmp_path):
    text = _long_text("Secret private content")
    path = tmp_path / "brief.docx"
    _write_docx(path, [text])

    asset = load_docx_asset(path, BRIEF)
    representation = repr(asset)

    assert "Secret private content" not in representation
    assert asset.content_hash in representation


def test_short_content_is_rejected(tmp_path):
    path = tmp_path / "asset.txt"
    path.write_text("too short", encoding="utf-8")

    with pytest.raises(ValueError):
        load_text_asset(path, BRIEF)
